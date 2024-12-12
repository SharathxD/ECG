import os
import tensorflow as tf
from datetime import datetime
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Google Generative AI
api_key = os.getenv("GEMINI_API_KEY")
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}
genai_model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

# Load the ECG model
ecg_model = tf.keras.models.load_model('ecg.h5')
chat = genai_model.start_chat(history=[])

# ECG classes
class_names = [
    'ECG Images of Myocardial Infarction Patients (240x12=2880)',
    'ECG Images of Patient that have History of MI (172x12=2064)',
    'ECG Images of Patient that have abnormal heartbeat (233x12=2796)',
    'Normal Person ECG Images (284x12=3408)'
]

# Page Config
st.set_page_config(page_title="CardioBuddy", page_icon="🫀", layout="wide")

# Custom CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&display=swap');

    body {
        font-family: 'Roboto', sans-serif;
        color: #333333;
        background-color: #f8f9fa;
    }
    .main {
        padding: 2rem;
        border-radius: 10px;
        background-color: #ffffff;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    h1, h2, h3 {
        color: #2c3e50;
        font-weight: 500;
    }
    .stButton>button {
        background-color: #3498db;
        color: white;
        font-weight: 500;
        border: none;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #2980b9;
        transform: translateY(-2px);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .feature-card {
        background-color: #ffffff;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 1rem;
        border-left: 5px solid #3498db;
        transition: all 0.3s ease;
    }
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
    }
    .feature-icon {
        font-size: 2rem;
        margin-right: 0.5rem;
        color: #3498db;
    }
    .stTextInput>div>div>input {
        border-color: #3498db;
        border-radius: 5px;
    }
    .stTextInput>div>div>input:focus {
        border-color: #2980b9;
        box-shadow: 0 0 0 1px #2980b9;
    }
    .stSelectbox>div>div>div {
        border-color: #3498db;
        border-radius: 5px;
    }
    .stSelectbox>div>div>div:focus {
        border-color: #2980b9;
        box-shadow: 0 0 0 1px #2980b9;
    }
    .chat-container {
        background-color: #ffffff;
        border-radius: 10px;
        padding: 20px;
        height: 60vh;
        display: flex;
        flex-direction: column;
        border: 1px solid #e0e0e0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .chat-messages {
        flex-grow: 1;
        overflow-y: auto;
        padding-right: 10px;
        margin-bottom: 20px;
    }
    .chat-message {
        margin-bottom: 15px;
        display: flex;
        align-items: flex-start;
    }
    .user-message {
        justify-content: flex-end;
    }
    .bot-message {
        justify-content: flex-start;
    }
    .message-content {
        max-width: 70%;
        padding: 10px 15px;
        border-radius: 18px;
        font-size: 14px;
        line-height: 1.4;
    }
    .user-message .message-content {
        background-color: #3498db;
        color: white;
        border-bottom-right-radius: 0;
    }
    .bot-message .message-content {
        background-color: #f0f3f6;
        color: #333333;
        border-bottom-left-radius: 0;
    }
    .chat-input {
        display: flex;
        margin-top: auto;
    }
    .chat-input input {
        flex-grow: 1;
        padding: 10px;
        border: 1px solid #3498db;
        border-radius: 20px;
        margin-right: 10px;
    }
    .chat-input button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 20px;
        padding: 10px 20px;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    .chat-input button:hover {
        background-color: #2980b9;
        transform: translateY(-2px);
    }
    .sidebar .sidebar-content {
        background-color: #2c3e50;
        color: white;
        padding: 20px;
    }
    .sidebar .sidebar-content .stRadio > label {
        color: white;
    }
    .stProgress > div > div > div {
        background-color: #3498db;
    }
</style>
""", unsafe_allow_html=True)

# Helper Functions
def get_gemini_response(input,history):
    prompt=f"""Initial Configuration

1. Role: Medical Assistant (CVD Specialist)
2. Knowledge Domain: Cardiovascular diseases, symptoms, diagnosis, treatment, and prevention.
3. Conversational Scope: Patient consultation, medical inquiry, and education.
Respond as a medical professional specializing in cardiovascular diseases. Address the user's query with accurate and relevant information. Provide clear explanations, diagnosis considerations, and treatment options. Limit responses to medical aspects only.

User Query: {input}

Generate a comprehensive response following these guidelines:

Response Guidelines

1. Accuracy: Ensure responses are evidence-based and up-to-date.
2. Clarity: Use simple, non-technical language.
3. Relevance: Focus on cardiovascular diseases.
4. Professionalism: Maintain a neutral, empathetic tone.
5. Contextual understanding: Consider patient history, symptoms, and medical context.

Output Requirements

1. Format: Structured response with headings and bullet points.
2. Length: Concise, ideally 200-300 words.
3. Tone: Professional, empathetic.

Example Response

For user query: 'What are symptoms of heart failure?'

Response:

Symptoms
- Shortness of breath (dyspnea)
- Fatigue
- Swelling in legs, ankles, and feet (edema)
- Rapid or irregular heartbeat
- Coughing up pink, frothy mucus

Medical Considerations
If experiencing symptoms, consult a healthcare provider for proper diagnosis and treatment.You are an AI assistant . Here's the conversation so far:\n\n
            {''.join([f'{role}:{text}'for role, text in history])}\n\n
            now, Here's the user's new query:{input}\n\n
            Please provide a comprehensive and informative response"""
    response=chat.send_message(prompt,stream=False)
    return response

def load_and_prep_image(image, img_shape=224):
    img = tf.convert_to_tensor(image)
    img = tf.image.decode_image(tf.io.encode_jpeg(img), channels=3)
    img = tf.image.resize(img, size=[img_shape, img_shape])
    img = img / 255.0
    return img

def pred_and_plot(model, image, class_names):
    img = load_and_prep_image(image)
    pred = model.predict(tf.expand_dims(img, axis=0))
    if len(pred[0]) > 1:
        pred_class = class_names[pred.argmax()]
    else:
        pred_class = class_names[int(tf.round(pred)[0][0])]
    return pred_class

def generate_ecg_details(ecg_image,lang):
    prompt = f"""
  1. Role: Pharmacy Assistant
2. Knowledge Domain: Pharmaceuticals, pharmacology, and prescription analysis
3. Conversational Scope: Prescription explanation and guidance

User Query
Analyze the attached prescription image and provide a detailed report in the {lang} only:

Image: {ecg_image}

Extract the following information from the prescription image:

1. Medication name
2. Dosage
3. Frequency
4. Duration
5. Special instructions

Provide a comprehensive response explaining:

1. Medication purpose and benefits
2. Dosage and frequency rationale
3. Potential side effects and interactions
4. Important precautions and warnings
5. Follow-up and monitoring recommendations
    
"""
    chat_session = genai_model.start_chat(history=[])
    response = chat_session.send_message([prompt])
    return response.text

def create_doc(report_text, ecg_image):
    doc = Document()
    doc.add_heading('ECG ANALYSIS REPORT', 0)
    
    for line in report_text.split("\n"):
        if line.strip():
            if line.startswith("**") and line.endswith("**"):
                doc.add_heading(line.strip("**"), level=1)
            elif line.startswith("-"):
                doc.add_paragraph(line.strip(), style="List Bullet")
            else:
                doc.add_paragraph(line.strip())

    doc.add_heading('ECG Tracing:', level=1)
    image_stream = BytesIO(ecg_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Sidebar Navigation
with st.sidebar:
    #st.image("NAV", width=100)
    st.title("CardioBuddy")
    nav = st.radio("Navigation", ["Home", "ECG Analysis", "Chatbot", "Gen AI"])

# Home Page
if nav == "Home":
    st.markdown("<h1 style='text-align: center;'>🫀 Welcome to CardioBuddy</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 1.2em;'>Your AI-powered cardiology assistant</p>", unsafe_allow_html=True)
    
    st.write("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class='feature-card'>
            <h3><span class='feature-icon'>🔬</span> ECG Analysis</h3>
            <p>Upload ECG images for instant AI-powered analysis and classification.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class='feature-card'>
            <h3><span class='feature-icon'>🤖</span> AI Chatbot</h3>
            <p>Consult our AI for healthcare-related queries and get evidence-based responses.</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class='feature-card'>
            <h3><span class='feature-icon'>🧠</span> Generative AI Reports</h3>
            <p>Generate comprehensive ECG reports using cutting-edge AI technology.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class='feature-card'>
            <h3><span class='feature-icon'>📊</span> Downloadable Reports</h3>
            <p>Create and download detailed ECG analysis reports in DOCX format.</p>
        </div>
        """, unsafe_allow_html=True)

# ECG Analysis Page
elif nav == "ECG Analysis":
    st.markdown("<h1 style='text-align: center;'>🔬 ECG Analysis</h1>", unsafe_allow_html=True)
    st.write("Upload an ECG image for instant AI-powered analysis and classification.")
    
    uploaded_image = st.file_uploader("Upload an ECG Image", type=["jpg", "jpeg", "png"])
    if uploaded_image is not None:
        image = Image.open(uploaded_image)
        st.image(image, caption="Uploaded ECG Image", use_container_width=True)
        with st.spinner("Analyzing ECG..."):
            pred_class = pred_and_plot(ecg_model, image, class_names)
        st.success(f"Analysis Result: {pred_class}")
        
        st.write("---")
        
        st.markdown("<h3>Interpretation</h3>", unsafe_allow_html=True)
        st.write("Based on the AI analysis, here's a brief interpretation of the ECG:")
        
        if "Myocardial Infarction" in pred_class:
            st.warning("This ECG shows signs consistent with a Myocardial Infarction (heart attack). Immediate medical attention is crucial.")
        elif "History of MI" in pred_class:
            st.info("This ECG indicates a history of Myocardial Infarction. Regular follow-ups and cardiac care are important.")
        elif "abnormal heartbeat" in pred_class:
            st.warning("This ECG shows an abnormal heartbeat pattern. Further investigation by a cardiologist is recommended.")
        else:
            st.success("This ECG appears to be within normal limits. However, always consult with a healthcare professional for a comprehensive evaluation.")
        
        st.write("Remember, AI analysis is a tool to assist medical professionals and should not replace expert medical advice.")

# Chatbot Page
elif nav == "Chatbot":
    st.markdown("<h1 style='text-align: center;'>🤖 AI Chatbot</h1>", unsafe_allow_html=True)
    st.write("Ask anything or consult the chatbot for healthcare-related queries.")

    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    
    # Chat messages
    st.markdown('<div class="chat-messages">', unsafe_allow_html=True)
    for role, text in st.session_state['chat_history']:
        if role == "User":
            st.markdown(f'<div class="chat-message user-message"><div class="message-content">{text}</div></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-message bot-message"><div class="message-content">{text}</div></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Chat input
    st.markdown('<div class="chat-input">', unsafe_allow_html=True)
    user_input = st.text_input("Type your message...", key="chat_input")
    send_button = st.button("Send")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    if send_button and user_input:
        st.session_state['chat_history'].append(("User", user_input))
        with st.spinner("Generating response..."):
            response = get_gemini_response(user_input, st.session_state['chat_history'])
        for chunk in response:
            st.session_state['chat_history'].append(("Bot", chunk.text))
        

# Gen AI Page
elif nav == "Gen AI":
    st.markdown("<h1 style='text-align: center;'>🧠 Generate prescription Report with AI</h1>", unsafe_allow_html=True)
    st.write("Upload an prescription image and select a language to generate a comprehensive report.")
    
    languages = [
        "English", "Assamese", "Bengali", "Bodo", "Dogri", "Gujarati", "Hindi", "Kannada", 
        "Kashmiri", "Konkani", "Maithili", "Malayalam", "Manipuri", "Marathi", 
        "Nepali", "Odia", "Punjabi", "Sanskrit", "Santali", "Sindhi", "Tamil", 
        "Telugu", "Urdu"
    ]
    lang = st.selectbox("Select Language for Report", options=languages, index=0)

    ecg_image = st.file_uploader("Upload ECG Image for Analysis", type=["png", "jpg", "jpeg"])
    
    if ecg_image is not None:
        st.image(ecg_image, caption="Uploaded ECG Image", use_container_width=True)
        
        if st.button("Generate ECG Report"):
            with st.spinner("Analyzing ECG image and generating report..."):
                ecg_details = generate_ecg_details(ecg_image, lang)
            
            st.success("Report generated successfully!")
            st.markdown("<h2 style='text-align: center;'>Generated ECG Report</h2>", unsafe_allow_html=True)
            st.markdown(ecg_details)
            
            doc_file_stream = create_doc(ecg_details, ecg_image)
            st.download_button(
                label="Download ECG Report",
                data=doc_file_stream,
                file_name=f"ECG_Report_{lang}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.write("---")
st.markdown("<p style='text-align: center;'>© 2024 CardioBuddy. All rights reserved.</p>", unsafe_allow_html=True)

